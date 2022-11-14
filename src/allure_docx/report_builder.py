import os
import warnings
import re
import shutil
import subprocess
import json
import sys

from os import listdir
from os.path import join, isfile
from time import ctime
from datetime import timedelta, datetime

from docx.shared import Mm, Cm
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import piechart


def _format_argval(argval):
    """Remove newlines and limit max length

    From Allure-pytest logger (formats argument in the CLI live logs).
    Consider using the same function."""

    max_arg_length = 100
    argval = argval.replace("\n", " ")
    if len(argval) > max_arg_length:
        argval = argval[:3] + " ... " + argval[-max_arg_length:]
    return argval


class ReportBuilder:
    """
    Builder to create a report from a given ReportConfig Object.
    """

    def __init__(self, config):
        self.indent = 6
        self.document = Document(config['template_path'])
        self.config = config

        self.session = {
            "allure_dir": config['allure_dir'],
            "start": None,
            "stop": None,
            "results": {
                "broken": 0,
                "failed": 0,
                "skipped": 0,
                "passed": 0,
            },
            "results_relative": {
                "broken": 0,
                "failed": 0,
                "skipped": 0,
                "passed": 0,
            },
            "total": 0,
        }

        self.sorted_results = None
        self._build_data()
        self._create_pie_chart()
        self._create_report()

    def save_report(self, output):
        """
        Save report to given output path as docx.
        """
        self.document.save(output)

    def save_report_to_pdf(self, output):
        """
        Save report to given output path as pdf. Tries officetopdf or soffice.
        """
        officetopdf = shutil.which("OfficeToPDF")
        soffice = shutil.which("soffice")

        docx_filename = f"{os.path.dirname(output)}/__temp.docx"
        self.save_report(docx_filename)

        if officetopdf is not None:
            print("Found OfficeToPDF, using it. Make sure you have MS Word installed.")
            proc = subprocess.run(
                [officetopdf, "/bookmarks", "/print", docx_filename, output],
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                check=False
            )
            print(proc.stdout.decode())
            sys.exit(proc.returncode)
        elif soffice is not None:
            result_dir = os.path.dirname(output)
            doc_path = output
            subprocess.call(["soffice", "--convert-to", "pdf", "--outdir", result_dir, doc_path])
        else:
            print("Could not find neither OfficeToPDF nor soffice. Not generating PDF.")

        os.remove(docx_filename)

    def _build_data(self):
        def process_steps(node):
            if "start" in node:
                if "start" not in self.session:
                    self.session["start"] = node["start"]
                elif self.session["start"] is None:
                    self.session["start"] = node["start"]
                elif node["start"] < self.session["start"]:
                    self.session["start"] = node["start"]

            if "stop" in node:
                if "stop" not in self.session:
                    self.session["stop"] = node["stop"]
                elif self.session["stop"] is None:
                    self.session["stop"] = node["stop"]
                elif node["stop"] > self.session["stop"]:
                    self.session["stop"] = node["stop"]

            if "steps" in node:
                for step in node["steps"]:
                    process_steps(step)

        def get_sorting_key(d):
            classification = {"broken": 0, "failed": 1, "skipped": 2, "passed": 3}
            return f"{classification[d['status']]}-{d['name']}"

        allure_dir = self.config['allure_dir']

        json_results = [f for f in listdir(allure_dir) if isfile(join(allure_dir, f)) and "result" in f]
        json_containers = [f for f in listdir(allure_dir) if isfile(join(allure_dir, f)) and "container" in f]

        data_containers = []
        for file in json_containers:
            with open(join(allure_dir, file), encoding="utf-8") as f:
                container = json.load(f)
                data_containers.append(container)

        data_results = []
        for file in json_results:
            with open(join(allure_dir, file), encoding="utf-8") as f:
                result = json.load(f)
                result["_lastmodified"] = os.path.getmtime(join(allure_dir, file))

                skip = False
                for previous_item in list(data_results):  # copy
                    if previous_item["name"] == result["name"]:
                        if previous_item["_lastmodified"] > result["_lastmodified"]:
                            skip = True
                        else:
                            data_results.remove(previous_item)
                        break
                if skip:
                    continue
                data_results.append(result)

        for result in data_results:
            process_steps(result)
            self.session["total"] += 1
            self.session["results"][result["status"]] += 1

            result["parents"] = []
            for container in data_containers:
                if "children" not in container:
                    continue
                if result["uuid"] in container["children"]:
                    result["parents"].append(container)
                    if "befores" in container:
                        for before in container["befores"]:
                            process_steps(before)
                    if "afters" in container:
                        for after in container["afters"]:
                            process_steps(after)

        if self.session["total"] == 0:
            warnings.warn("No test result files were found!")

        self.sorted_results = sorted(data_results, key=get_sorting_key)

        if self.session["start"] is not None:
            self.session["duration"] = str(timedelta(seconds=(self.session["stop"] - self.session["start"]) / 1000.0))
            self.session["start"] = ctime(self.session["start"] / 1000.0)
            self.session["stop"] = ctime(self.session["stop"] / 1000.0)
        else:
            self.session["duration"] = "Not available"
            self.session["start"] = "Not available"
            self.session["stop"] = "Not available"

        for item in self.session["results"]:
            if self.session["total"] > 0:
                self.session["results_relative"][item] = "{:.2f}%".format(
                    100 * self.session["results"][item] / self.session["total"])
            else:
                self.session["results_relative"][item] = "Not available"

    def _create_pie_chart(self):
        img_file = os.path.join(self.session["allure_dir"], "pie.png")
        self.session["pie_chart_source"] = img_file
        piechart.create_piechart(self.session["results"], img_file)

    def _create_toc(self):
        # Snippet from:
        # https://github.com/python-openxml/python-docx/issues/36
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        fld_char = OxmlElement("w:fldChar")  # creates a new element
        fld_char.set(qn("w:fldCharType"), "begin")  # sets attribute on element
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")  # sets attribute on element
        instr_text.text = 'TOC \\o "1-1" \\h \\z'  # change 1-3 depending on heading levels you need

        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "separate")
        fld_char3 = OxmlElement("w:t")
        fld_char3.text = "Right-click to update field."
        fld_char2.append(fld_char3)

        fld_char4 = OxmlElement("w:fldChar")
        fld_char4.set(qn("w:fldCharType"), "end")

        r_element = run._r
        r_element.append(fld_char)
        r_element.append(instr_text)
        r_element.append(fld_char2)
        r_element.append(fld_char4)
        # p_element = paragraph._p

    def _print_attachments(self, item):
        if "attachments" in item:
            for attachment in item["attachments"]:
                self.document.add_paragraph(f"[Attachment] {attachment['name']}", style="Step")
                if "image" in attachment["type"]:
                    self.document.add_picture(
                        os.path.join(self.session["allure_dir"], attachment["source"]),
                        width=Mm(100),
                    )
                    self.document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _print_steps(self, parent_step, config_info, indent=0):
        indent_str = indent * self.indent * " "
        if "steps" in parent_step:
            for step in parent_step["steps"]:
                if step["status"] in ["failed", "broken"]:
                    step_style = "Step Failed"
                else:
                    step_style = "Step"
                self.document.add_paragraph(f"{indent_str}> {step['name']}", style=step_style)
                if "parameters" in config_info and "parameters" in step:
                    for params in step["parameters"]:
                        paragraph = self.document.add_paragraph(f"{indent_str}    ", style="Step Param Parag")
                        paragraph.add_run(
                            f"{params['name']} = {_format_argval(params['value'])}",
                            style="Step Param",
                        )
                if "details" in config_info and "statusDetails" in step and len(step["statusDetails"]) != 0:
                    if "message" in step["statusDetails"] and len(step["statusDetails"]["message"]) != 0:
                        self.document.add_paragraph(step["statusDetails"]["message"], style=step_style)

                    if "trace" in config_info and "trace" in step["statusDetails"] and len(
                            step["statusDetails"]["trace"]) != 0:
                        table = self.document.add_table(rows=1, cols=1, style="Trace table")
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].add_paragraph(step["statusDetails"]["trace"] + "\n", style="Code")
                if "attachments" in config_info:
                    self._print_attachments(step)
                self._print_steps(step, config_info, indent + 1)

    def _add_field(self, run, field):
        def create_attribute(element, name, value):
            element.set(qn(name), value)

        def create_element(name):
            return OxmlElement(name)

        fld_char1 = create_element('w:fldChar')
        create_attribute(fld_char1, 'w:fldCharType', 'begin')

        instr_text = create_element('w:instrText')
        create_attribute(instr_text, 'xml:space', 'preserve')
        instr_text.text = field

        fld_char2 = create_element('w:fldChar')
        create_attribute(fld_char2, 'w:fldCharType', 'end')

        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)

    def _create_footer(self, footer):
        footer.paragraphs[0].text += datetime.today().strftime('%Y-%m-%d')
        footer.paragraphs[0].text += "\t\t"
        footer_run = footer.paragraphs[0].add_run()
        self._add_field(footer_run, field="PAGE")

    def _delete_paragraph(self, paragraph):
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
        p_element._p = p_element._element = None

    def _create_header(self, header, details=False):
        htable = header.add_table(1, 2, Cm(16))
        htable.style = "header table"
        htab_cells = htable.rows[0].cells

        if self.config['logo']['path'] is not None:
            ht1 = htab_cells[1].add_paragraph()
            ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            kh = ht1.add_run()
            kh.add_picture(self.config['logo']['path'], width=Cm(5))

        if details:
            header_text = "Test Report"
            header_text += "\n" + self.config['cover']['title']
            if 'device_under_test' in self.config['details']:
                header_text += "\n" + self.config['details']['device_under_test']
            htab_cells[0].add_paragraph(header_text)
            self._delete_paragraph(header.paragraphs[0])
            self._delete_paragraph(htab_cells[0].paragraphs[0])
            self._delete_paragraph(htab_cells[1].paragraphs[0])
            header.add_paragraph("")

    def _print_cover(self):
        header = self.document.sections[0].header
        self._create_header(header)

        self._delete_paragraph(self.document.paragraphs[-0])
        if 'company_name' in self.config['cover']:
            self.document.add_paragraph("\n" + self.config['cover']['company_name'], style="company")
        self.document.add_paragraph("\n\n\n\nTest Report", style="Title")
        subtitle = self.config['cover']['title']
        if 'device_under_test' in self.config['details']:
            subtitle += "\n" + self.config['details']['device_under_test']
        self.document.add_paragraph(subtitle, style="Subtitle")
        self.document.add_paragraph("\n" + datetime.today().strftime('%Y-%m-%d'), style="heading 2")

    def _print_details(self):
        if 'details' in self.config and len(self.config['details']) > 0:
            self.document.add_paragraph("Test Details", style="Heading 1")

            i = 0
            detail_table = self.document.add_table(rows=len(self.config['details']), cols=2, style="Label table")
            for detail in self.config['details'].items():
                detail_table.rows[i].cells[0].paragraphs[-1].clear().add_run(detail[0].replace("_", " ").capitalize())
                detail_table.rows[i].cells[1].paragraphs[-1].clear().add_run(re.sub(r";\s*", "\n", detail[1]))
                i += 1

    def _print_session_summary(self):
        self.document.add_paragraph("Test Session Summary", style="Heading 1")

        table = self.document.add_table(rows=1, cols=2)
        summary_cell = table.rows[0].cells[0]
        summary_cell.add_paragraph(
            f"Start: {self.session['start']}\nEnd: {self.session['stop']}\nDuration: {self.session['duration']}"
        )
        self._delete_paragraph(summary_cell.paragraphs[0])

        results_strs = []
        for item in self.session["results"]:
            results_strs.append(f"{item}: {self.session['results'][item]} ({self.session['results_relative'][item]})")
        summary_cell.add_paragraph("\n".join(results_strs))

        pie_chart_cell = table.rows[0].cells[1]
        paragraph = pie_chart_cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(self.session["pie_chart_source"], width=Mm(75))

    def _print_test(self, test):
        # config elements for the specific status of this test
        config_info = self.config["info"][test["status"]]
        config_labels = self.config["labels"][test["status"]]

        self.document.add_paragraph(f"{test['name']}  [ {test['status']} ]", style=f"Heading {test['status']}")

        table = None
        added_table = False
        if "duration" in config_info:
            duration = test["stop"] - test["start"]
            duration_unit = "ms"
            if duration > 1000:
                duration_unit = "s"
                duration = duration / 1000
                if duration > 60:
                    duration_unit = "min"
                    duration = duration / 60

            table = self.document.add_table(rows=1, cols=2, style="Label table")
            table.rows[0].cells[0].paragraphs[-1].clear().add_run("Duration")
            table.rows[0].cells[1].paragraphs[-1].clear().add_run(str(duration) + duration_unit)
            added_table = True

        # add labels to table
        for label_name in config_labels:
            if not added_table:
                table = self.document.add_table(rows=0, cols=2, style="Label table")
                added_table = True
            iterator = iter(label for label in test["labels"] if label["name"].lower() == label_name)
            label = next(iterator, None)
            if label is not None:
                row = table.add_row()
                row.cells[0].paragraphs[-1].clear().add_run(label_name.capitalize())
                while label is not None:
                    row.cells[1].add_paragraph(label["value"])
                    label = next(iterator, None)
                self._delete_paragraph(row.cells[1].paragraphs[0])

        if table is not None:
            table.columns[0].width = Cm(4)
            for cell in table.columns[0].cells:
                cell.width = Cm(4)
            table.columns[1].width = Cm(12)
            for cell in table.columns[1].cells:
                cell.width = Cm(12)
            self.document.add_paragraph()

        if "description" in config_info:
            self.document.add_heading("Description", level=2)
            if "description" in test and len(test["description"]) != 0:
                self.document.add_paragraph(test["description"])
            else:
                self.document.add_paragraph("No description available.")

        if "parameters" in config_info and "parameters" in test and len(test["parameters"]) != 0:
            self.document.add_heading("Parameters", level=2)
            for p in test["parameters"]:
                self.document.add_paragraph(f"{p['name']}: {p['value']}", style="Step")

        if (
                "details" in config_info
                and "statusDetails" in test
                and len(test["statusDetails"]) != 0
                and (
                "message" in test["statusDetails"]
                and len(test["statusDetails"]["message"]) != 0
                or "trace" in config_info
                and "trace" in test["statusDetails"]
        )
        ):
            self.document.add_heading("Details", level=2)
            if "message" in test["statusDetails"]:
                self.document.add_paragraph(test["statusDetails"]["message"], style=None)
            if "trace" in config_info and "trace" in test["statusDetails"]:
                table = self.document.add_table(rows=1, cols=1, style="Trace table")
                hdr_cells = table.rows[0].cells
                hdr_cells[0].add_paragraph(test["statusDetails"]["trace"] + "\n", style="Code")

        if "links" in config_info and "links" in test and len(test["links"]) != 0:
            self.document.add_heading("Links", level=2)
            for link in test["links"]:
                if "name" in link and "url" in link:
                    self.document.add_paragraph(f"{link['name']}: {link['url']}")
                else:
                    print("WARNING: A link was provided without name or url and will not be printed.")

        if "setup" in config_info:
            self.document.add_heading("Test Setup", level=2)
            for parent in test["parents"]:
                if "befores" in parent:
                    for before in parent["befores"]:
                        self.document.add_paragraph(f"[Fixture] {before['name']}", style="Step")
                        self._print_attachments(before)
                        self._print_steps(before, config_info, 1)
            if self.document.paragraphs[-1].text == "Test Setup":
                self._delete_paragraph(self.document.paragraphs[-1])

        if "body" in config_info:
            self.document.add_heading("Test Body", level=2)
            self._print_attachments(test)
            self._print_steps(test, config_info)
            if self.document.paragraphs[-1].text == "Test Body":
                self._delete_paragraph(self.document.paragraphs[-1])

        if "teardown" in config_info:
            self.document.add_heading("Test Teardown", level=2)
            for parent in test["parents"]:
                if "afters" in parent:
                    for after in parent["afters"]:
                        self.document.add_paragraph("[Fixture] {after['name']}", style="Step")
                        self._print_attachments(after)
                        self._print_steps(after, config_info, 1)
            if self.document.paragraphs[-1].text == "Test Teardown":
                self._delete_paragraph(self.document.paragraphs[-1])

    def _create_report(self):
        if not self.sorted_results:
            self.document.add_paragraph("No test result files were found.")
            self.document.save_report(self.config['output_filename'])
            return

        self._print_cover()
        self.document.add_section()

        footer = self.document.sections[1].footer
        footer.is_linked_to_previous = False
        self._create_footer(footer)
        header = self.document.sections[1].header
        header.is_linked_to_previous = False
        self._create_header(header, True)

        self._print_details()
        self._print_session_summary()

        self.document.add_page_break()

        # print tests
        for test in self.sorted_results:
            self._print_test(test)
